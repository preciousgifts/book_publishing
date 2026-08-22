import io
import re
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from html.parser import HTMLParser
from config import get_trim_spec

PAGE_TITLES = {
    'title_page': 'Title Page',
    'copyright_page': 'Copyright Page',
    'dedication': 'Dedication',
    'epigraph': 'Epigraph',
    'table_of_contents': 'Table of Contents',
    'foreword': 'Foreword',
    'preface': 'Preface',
    'acknowledgments': 'Acknowledgments',
    'introduction': 'Introduction',
    'appendix': 'Appendix',
    'glossary': 'Glossary',
    'bibliography': 'Endnotes & Bibliography',
    'index': 'Index',
    'about_author': 'About the Author',
    'also_by_author': 'Also By the Author',
    'discussion_questions': 'Discussion Questions',
    'call_to_action': 'Reader Review & Connect'
}

class ParagraphBuilder(HTMLParser):
    def __init__(self, doc: Document, initial_paragraph, spec: dict, default_align=WD_ALIGN_PARAGRAPH.JUSTIFY, default_font_size=None):
        super().__init__()
        self.doc = doc
        self.spec = spec
        self.current_paragraph = initial_paragraph
        self.default_align = default_align
        self.default_font_size = default_font_size
        self.active_styles = set()
        self.has_content = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attr_dict = dict(attrs)

        if tag == "p":
            if self.has_content or len(self.current_paragraph.runs) > 0 or self.current_paragraph.text:
                self.current_paragraph = self.doc.add_paragraph()
            
            align_style = attr_dict.get("style", "")
            if "text-align: center" in align_style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "text-align: right" in align_style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif "text-align: left" in align_style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                self.current_paragraph.alignment = self.default_align

            if "text-indent" in align_style or "indent-first-line" in attr_dict.get("class", ""):
                self.current_paragraph.paragraph_format.first_line_indent = Inches(0.25)

        elif tag == "h1":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                self.current_paragraph.style = "Heading 1"
            except Exception:
                pass
        elif tag == "h2":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                self.current_paragraph.style = "Heading 2"
            except Exception:
                pass
        elif tag == "h3":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                self.current_paragraph.style = "Heading 3"
            except Exception:
                pass
        elif tag == "blockquote":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            self.current_paragraph.paragraph_format.right_indent = Inches(0.5)
            self.active_styles.add("italic")
        elif tag in ("strong", "b"):
            self.active_styles.add("bold")
        elif tag in ("em", "i"):
            self.active_styles.add("italic")
        elif tag == "br":
            self.current_paragraph.add_run().add_break()
            self.has_content = True
        elif tag in ("img", "figure"):
            src = attr_dict.get("src") or attr_dict.get("data-src")
            caption = attr_dict.get("data-caption") or attr_dict.get("alt")
            if src and "data:image/" in src and ";base64," in src:
                try:
                    base64_data = src.split(";base64,")[1]
                    img_bytes = base64.b64decode(base64_data)
                    img_p = self.doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_p.add_run()
                    max_width = Inches(min(5.0, self.spec["width"] * 0.75))
                    run.add_picture(io.BytesIO(img_bytes), width=max_width)

                    if caption:
                        cap_p = self.doc.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(f"Figure: {caption}")
                        cap_run.italic = True
                        cap_run.font.size = Pt(9)
                except Exception as e:
                    pass

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("strong", "b"):
            self.active_styles.discard("bold")
        elif tag in ("em", "i", "blockquote"):
            self.active_styles.discard("italic")

    def handle_data(self, data):
        if not data:
            return
        clean_text = re.sub(r'^#{1,6}\s+', '', data)
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
        clean_text = re.sub(r'__(.*?)__', r'\1', clean_text)
        clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
        clean_text = re.sub(r'_(.*?)_', r'\1', clean_text)
        
        run = self.current_paragraph.add_run(clean_text)
        self.has_content = True
        if "bold" in self.active_styles:
            run.bold = True
        if "italic" in self.active_styles:
            run.italic = True
        if self.default_font_size:
            run.font.size = Pt(self.default_font_size)


def render_matter_page_to_docx(doc: Document, page: dict, spec: dict):
    page_type = page.get("pageType", "")
    content = page.get("content", "")
    status = page.get("status", "NOT_GENERATED")
    title_label = PAGE_TITLES.get(page_type, page_type.replace('_', ' ').title())

    # Heading unless Title Page
    if page_type != 'title_page':
        h_p = doc.add_paragraph()
        h_p.style = "Heading 1"
        if page_type in ('copyright_page', 'dedication', 'epigraph'):
            h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if page_type != 'copyright_page' else WD_ALIGN_PARAGRAPH.LEFT
        else:
            h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_p.add_run(title_label)

    if not content or content.strip() == "" or status == "NOT_GENERATED":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{title_label} - Content not yet generated for this page]")
        run.italic = True
        run.font.size = Pt(10)
        return

    # Specific formatting per pageType
    if page_type == 'title_page':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder = ParagraphBuilder(doc, p, spec, default_align=WD_ALIGN_PARAGRAPH.CENTER)
        builder.feed(content)

    elif page_type == 'copyright_page':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        builder = ParagraphBuilder(doc, p, spec, default_align=WD_ALIGN_PARAGRAPH.LEFT, default_font_size=9)
        builder.feed(content)

    elif page_type in ('dedication', 'epigraph'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder = ParagraphBuilder(doc, p, spec, default_align=WD_ALIGN_PARAGRAPH.CENTER)
        builder.feed(content)

    else:
        # Standard prose pages (Foreword, Preface, Ack, Intro, Appendix, Glossary, Bibliography, Index, About Author, etc.)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        builder = ParagraphBuilder(doc, p, spec, default_align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        builder.feed(content)


def build_docx(paragraphs_data: list, trim_size: str, matter_pages: list = None, book_title: str = "") -> Document:
    """
    Builds a word document with justified alignment, figure support, front/back matter pages, and mirror margins.
    """
    doc = Document()
    spec = get_trim_spec(trim_size)
    section = doc.sections[0]
    
    # Page setup
    section.page_width = Inches(spec["width"])
    section.page_height = Inches(spec["height"])
    
    # Margin settings
    section.top_margin = Inches(spec["margins"]["top"])
    section.bottom_margin = Inches(spec["margins"]["bottom"])
    section.left_margin = Inches(spec["margins"]["inside"])   # inside
    section.right_margin = Inches(spec["margins"]["outside"])  # outside
    
    # Word XML Mirror Margins Injection
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    type_elm = parse_xml(r'<w:mirrorMargins xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    sectPr.append(type_elm)

    matter_pages = matter_pages or []
    front_matter = [m for m in matter_pages if m.get("section") == "FRONT"]
    back_matter = [m for m in matter_pages if m.get("section") == "BACK"]

    front_matter.sort(key=lambda x: x.get("order", 0))
    back_matter.sort(key=lambda x: x.get("order", 0))

    # 1. Render Front Matter Pages
    for idx, page in enumerate(front_matter):
        render_matter_page_to_docx(doc, page, spec)
        doc.add_page_break()

    # 2. Render Chapters
    if not paragraphs_data:
        if not front_matter and not back_matter:
            doc.add_paragraph("")
    else:
        current_chapter = None
        for p_data in paragraphs_data:
            ch_idx = p_data.get("chapterIndex", 0)
            if current_chapter is None or ch_idx != current_chapter:
                if current_chapter is not None:
                    doc.add_page_break()
                current_chapter = ch_idx
                ch_p = doc.add_paragraph()
                ch_p.style = "Heading 1"
                ch_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                ch_p.add_run(f"Chapter {current_chapter + 1}")
                active_p = doc.add_paragraph()
                active_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                active_p = doc.add_paragraph()
                active_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            builder = ParagraphBuilder(doc, active_p, spec)
            builder.feed(p_data.get("formattedHtml", ""))

    # 3. Render Back Matter Pages
    for page in back_matter:
        doc.add_page_break()
        render_matter_page_to_docx(doc, page, spec)

    # Cleanup trailing/empty paragraphs
    for p in list(doc.paragraphs):
        if len(p.runs) == 0 and not p.text:
            p_element = p._p
            p_element.getparent().remove(p_element)
            
    return doc
