import pytest
import io
import docx
from unittest.mock import AsyncMock, patch
from fastapi.testclient import TestClient

from main import app
from pipeline.manuscript_parser import parse_manuscript_file

client = TestClient(app)

@pytest.mark.asyncio
@patch("pipeline.manuscript_parser.AuditorAgent")
async def test_parse_manuscript_file_text(mock_auditor_class):
    mock_auditor = AsyncMock()
    mock_auditor.audit_manuscript.return_value = {
        "healthReport": "### Health Report\nPacing looks great.",
        "toc": [
            {"chapterNumber": 1, "title": "Chapter 1: Introduction", "summary": "Start of text", "subtopics": []}
        ]
    }
    mock_auditor_class.return_value = mock_auditor

    text_content = """Chapter 1: Introduction
This is the first paragraph.
This is the second paragraph.

Chapter 2: Scaling Up
More text for the second chapter.
"""
    result = await parse_manuscript_file(
        file_bytes=text_content.encode("utf-8"),
        filename="manuscript.txt",
        title="My Custom Text Book",
        genre="non-fiction",
        locale="en-US"
    )

    assert result["healthReport"] == "### Health Report\nPacing looks great."
    assert result["paragraphs"][0]["chapterIndex"] == 0
    assert result["paragraphs"][2]["chapterIndex"] == 1
    assert result["paragraphs"][0]["rawContent"] == "This is the first paragraph."

@pytest.mark.asyncio
@patch("pipeline.manuscript_parser.AuditorAgent")
async def test_parse_manuscript_markdown_headings(mock_auditor_class):
    mock_auditor = AsyncMock()
    mock_auditor.audit_manuscript.return_value = {"healthReport": "OK", "toc": []}
    mock_auditor_class.return_value = mock_auditor

    md_content = """# Introduction to Publishing
First paragraph of introduction.

## Deep Dive into AI
Prose discussing AI tools.

## Final Summary
Concluding thoughts.
"""
    result = await parse_manuscript_file(
        file_bytes=md_content.encode("utf-8"),
        filename="manuscript.md",
        title="Markdown Book",
        genre="non-fiction",
        locale="en-US"
    )

    assert result["paragraphs"][0]["chapterIndex"] == 0
    assert result["paragraphs"][1]["chapterIndex"] == 1
    assert result["paragraphs"][2]["chapterIndex"] == 2

@pytest.mark.asyncio
@patch("pipeline.manuscript_parser.AuditorAgent")
async def test_parse_manuscript_numbered_titles(mock_auditor_class):
    mock_auditor = AsyncMock()
    mock_auditor.audit_manuscript.return_value = {"healthReport": "OK", "toc": []}
    mock_auditor_class.return_value = mock_auditor

    txt_content = """1. Overview of the Strategy
Strategy overview details here.

2. Execution & Operations
Operational step by step text.
"""
    result = await parse_manuscript_file(
        file_bytes=txt_content.encode("utf-8"),
        filename="manuscript.txt",
        title="Numbered Book",
        genre="non-fiction",
        locale="en-US"
    )

    assert result["paragraphs"][0]["chapterIndex"] == 0
    assert result["paragraphs"][1]["chapterIndex"] == 1

@pytest.mark.asyncio
@patch("pipeline.manuscript_parser.AuditorAgent")
async def test_parse_manuscript_docx_with_styles(mock_auditor_class):
    mock_auditor = AsyncMock()
    mock_auditor.audit_manuscript.return_value = {"healthReport": "OK", "toc": []}
    mock_auditor_class.return_value = mock_auditor

    doc = docx.Document()
    doc.add_heading("Heading Chapter One", level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Bold statement here.")
    r1.bold = True
    
    doc.add_heading("Heading Chapter Two", level=1)
    p2 = doc.add_paragraph("Second chapter text.")

    bio = io.BytesIO()
    doc.save(bio)

    result = await parse_manuscript_file(
        file_bytes=bio.getvalue(),
        filename="manuscript.docx",
        title="Docx Book",
        genre="non-fiction",
        locale="en-US"
    )

    assert len(result["paragraphs"]) == 2
    assert result["paragraphs"][0]["chapterIndex"] == 0
    assert "<strong>Bold statement here.</strong>" in result["paragraphs"][0]["formattedHtml"]
    assert result["paragraphs"][1]["chapterIndex"] == 1

@pytest.mark.asyncio
@patch("main.parse_manuscript_file")
async def test_fastapi_parse_manuscript_endpoint(mock_parse_file):
    mock_parse_file.return_value = {
        "healthReport": "Audit report text",
        "toc": [{"chapterNumber": 1, "title": "Chapter 1", "summary": "...", "subtopics": []}],
        "paragraphs": [{"chapterIndex": 0, "paragraphIndex": 0, "rawContent": "...", "formattedHtml": "..."}]
    }

    response = client.post(
        "/internal/manuscript/parse",
        data={"title": "Uploaded Book", "genre": "fiction", "languageLocale": "en-US", "trimSize": "6x9"},
        files={"file": ("manuscript.txt", b"Chapter 1\nHello world", "text/plain")}
    )

    assert response.status_code == 200
    data = response.json()
    assert data["healthReport"] == "Audit report text"
    assert len(data["toc"]) == 1

@pytest.mark.asyncio
@patch("main.ArchitectAgent")
async def test_fastapi_adjust_outline_endpoint(mock_architect_class):
    mock_architect = AsyncMock()
    mock_architect.adjust_outline.return_value = {
        "toc": [
            {"chapterNumber": 1, "title": "Expanded Chapter 1", "summary": "...", "subtopics": []}
        ]
    }
    mock_architect_class.return_value = mock_architect

    payload = {
        "projectId": "3f8b9d3b-0bb5-4f35-86ff-94b12d5df835",
        "action": "expand",
        "targetChapterCount": 2,
        "feedback": "make it longer",
        "currentToC": [{"chapterNumber": 1, "title": "Chapter 1", "summary": "..."}]
    }

    response = client.post("/internal/swarm/adjust-outline", json=payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["tocData"][0]["title"] == "Expanded Chapter 1"
